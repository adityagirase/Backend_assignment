from typing import Dict, Any, BinaryIO
import PyPDF2
import docx
import json
import os

class DocumentProcessor:
    def process_document(self, file_obj: BinaryIO, filename: str) -> Dict[str, Any]:
        """Extract text content from various document formats"""
        _, ext = os.path.splitext(filename)
        ext = ext.lower()
        
        if ext == '.pdf':
            return self._process_pdf(file_obj)
        elif ext == '.docx':
            return self._process_docx(file_obj)
        elif ext == '.json':
            return self._process_json(file_obj)
        elif ext == '.txt':
            return self._process_txt(file_obj)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    
    def _process_pdf(self, file_obj):
        reader = PyPDF2.PdfReader(file_obj)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        
        return {
            "content": text,
            "metadata": {
                "page_count": len(reader.pages)
            }
        }
        
    def _process_docx(self, file_obj):
        doc = docx.Document(file_obj)
        text = "\n".join([para.text for para in doc.paragraphs])
        
        return {
            "content": text,
            "metadata": {
                "paragraph_count": len(doc.paragraphs)
            }
        }
    
    def _process_json(self, file_obj):
        data = json.load(file_obj)
        # For JSON, we'll store the raw JSON and also a text representation
        text_content = json.dumps(data, indent=2)
        
        return {
            "content": text_content,
            "raw_json": data,
            "metadata": {
                "is_structured": True
            }
        }
    
    def _process_txt(self, file_obj):
        content = file_obj.read().decode('utf-8')
        
        return {
            "content": content,
            "metadata": {
                "char_count": len(content)
            }
        }